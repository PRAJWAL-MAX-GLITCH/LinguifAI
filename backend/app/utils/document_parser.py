from __future__ import annotations

import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

# Maximum characters per chunk sent to the AI
CHUNK_SIZE = 1500


def parse_txt(file_path: Path) -> str:
    """Read plain text file as-is."""
    return file_path.read_text(encoding="utf-8", errors="ignore")


def parse_pdf(file_path: Path) -> str:
    """Extract text from PDF using pdfplumber, preserving page breaks."""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is not installed. Run: pip install pdfplumber")

    pages: List[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())

    if not pages:
        raise ValueError("No extractable text found in PDF.")
    return "\n\n".join(pages)


def parse_docx(file_path: Path) -> str:
    """Extract text from DOCX preserving paragraph structure."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise ValueError("No extractable text found in DOCX.")
    return "\n\n".join(paragraphs)


def extract_text(file_path: Path, extension: str) -> str:
    """
    Extract raw text from a document based on its file extension.

    Args:
        file_path: Absolute path to the file on disk.
        extension: Lowercased file extension (e.g. 'pdf', 'docx', 'txt').

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError if the extension is unsupported.
    """
    ext = extension.lower()
    if ext == "txt":
        return parse_txt(file_path)
    elif ext == "pdf":
        return parse_pdf(file_path)
    elif ext == "docx":
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: '{ext}'")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> List[str]:
    """
    Split text into sentence-aware chunks that respect the chunk_size limit.

    Splits on double newlines first (paragraphs), then on sentences
    to avoid cutting mid-thought.

    Args:
        text: The full extracted text.
        chunk_size: Maximum characters per chunk.

    Returns:
        A list of text chunks.
    """
    paragraphs = text.split("\n\n")
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(current) + len(para) + 2 <= chunk_size:
            current = f"{current}\n\n{para}".strip()
        else:
            if current:
                chunks.append(current)
            # If paragraph itself exceeds chunk_size, split by sentences
            if len(para) > chunk_size:
                sentences = para.replace(". ", ".|").split("|")
                temp = ""
                for sentence in sentences:
                    if len(temp) + len(sentence) + 1 <= chunk_size:
                        temp = f"{temp} {sentence}".strip()
                    else:
                        if temp:
                            chunks.append(temp)
                        temp = sentence
                if temp:
                    chunks.append(temp)
            else:
                current = para

    if current:
        chunks.append(current)

    logger.info("Chunked document into %d chunks (avg %d chars)", len(chunks), chunk_size)
    return chunks


def write_txt(translated_text: str, output_path: Path) -> None:
    """Write translated text to a plain text file."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(translated_text, encoding="utf-8")


def write_docx(translated_text: str, output_path: Path) -> None:
    """Write translated text to a DOCX file, preserving paragraph breaks."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx is not installed.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()
    for paragraph in translated_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(output_path)


def build_translated_output(
    translated_text: str, output_path: Path, extension: str
) -> None:
    """
    Persist the translated text in the same format as the source document.

    Args:
        translated_text: Full translated text content.
        output_path: Destination path for the output file.
        extension: File extension to determine output format.
    """
    ext = extension.lower()
    if ext == "txt":
        write_txt(translated_text, output_path)
    elif ext in ("pdf", "docx"):
        # Write DOCX for both PDF and DOCX originals (PDF generation requires additional deps)
        docx_path = output_path.with_suffix(".docx")
        write_docx(translated_text, docx_path)
        logger.info("PDF originals are output as DOCX: %s", docx_path)
    else:
        write_txt(translated_text, output_path)
